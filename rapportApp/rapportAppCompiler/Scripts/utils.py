from docx import Document
import subprocess
import sys
import os
import glob
import re
from Scripts.constants import OUTPUT_DIR


def getAllText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)


def showStructure(document):
    # Flag to indicate that the next paragraph is text under a heading
    print_next_paragraph = False

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            # Print the heading
            if paragraph.style.name == "Heading 1":
                print(f"|-- {paragraph.text}")
            elif paragraph.style.name == "Heading 2":
                indent = "    " * 1
                print(f"{indent}|-- {paragraph.text}")
            elif paragraph.style.name == "Heading 3":
                indent = "    " * 2
                print(f"{indent}|-- {paragraph.text}")
            elif paragraph.style.name == "Heading 4":
                indent = "    " * 3
                print(f"{indent}|--  {paragraph.text}")
            elif paragraph.style.name == "Heading 5":
                indent = "    " * 4
                print(f"{indent}|-- {paragraph.text}")

            # Set flag to true to print the next paragraph (if it's not a heading)
            print_next_paragraph = True
        else:
            # Check if the paragraph is text under a heading
            if print_next_paragraph:
                print(f"{indent}    {paragraph.text}")
                # Reset the flag
                print_next_paragraph = False


def run_pdflatex(filepath):
    # Save the current directory
    original_dir = os.getcwd()

    # Extract directory and filename from the provided path
    file_dir, filename = os.path.split(filepath)

    # Change to the directory where the .tex file is located
    os.chdir(file_dir)

    # Run the pdflatex command
    command = ["pdflatex", filename]
    process = subprocess.Popen(
        command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True
    )

    # Print the output as it is being generated
    for line in iter(process.stdout.readline, ""):
        print(line, end="")

    # Wait for the process to finish and get the exit code
    process.communicate()

    # Change back to the original directory
    os.chdir(original_dir)

    return process.returncode


def clean_files(
    directory,
    extensions=[
        ".toc",
        ".aux",
        ".xml",
        ".out",
        ".log",
        ".bcf",
        ".jpg",
        ".png",
        ".jpeg",
    ],
):
    if os.path.exists(directory):
        for ext in extensions:
            # Create a pattern for the current extension
            pattern = os.path.join(directory, f"*{ext}")

            # Find all files in the directory with the current extension
            files = glob.glob(pattern)

            # Delete each file
            for file in files:
                try:
                    os.remove(file)
                    print(f"Deleted: {file}")
                except OSError as e:
                    print(f"Error: {e.strerror} while deleting file {e.filename}")
    if directory == OUTPUT_DIR:
        os.startfile(OUTPUT_DIR)


def caption_validation(caption):
    # Check if the caption is empty
    if not caption.strip():
        return "Please add a caption to your image"

    # Pattern to match 'figure' or 'Figure' followed by a space and a number
    pattern = r"(?i)(figure\s+\d+)"

    # Remove the matched pattern from the caption
    processed_caption = re.sub(pattern, "", caption).strip()

    return processed_caption


def obj_type_from_caption(caption):
    # Check if the caption is empty
    print("\n\n checking caption", caption)
    if not caption.strip():
        return "Please add a caption."

    if is_image_caption(caption):
        # Pattern to match 'figure' or 'Figure' followed by a space and a number
        pattern = r"(?i)(figure\s+\d+)"

        # Remove the matched pattern from the caption
        processed_caption = re.sub(pattern, "", caption).strip()

        return processed_caption, "img"

    if is_table_caption(caption):
        # Pattern to match 'tableau' or 'Tableau' followed by a space and a number
        pattern = r"(?i)(tableau\s+\d+)"

        # Remove the matched pattern from the caption
        processed_caption = re.sub(pattern, "", caption).strip()

        return processed_caption, "table"


def is_image_caption(string):
    prefixes = ["figure", "Illustration"]
    # Remove leading whitespace and convert the string to lowercase
    string_lower = string.lstrip().lower()

    # Convert all prefixes to lowercase
    prefixes_lower = [prefix.lower() for prefix in prefixes]

    # Check if the string starts with any of the prefixes
    return any(string_lower.startswith(prefix) for prefix in prefixes_lower)


def is_table_caption(string):
    prefixes = ["table", "tableau"]
    # Remove leading whitespace and convert the string to lowercase
    string_lower = string.lstrip().lower()

    # Convert all prefixes to lowercase
    prefixes_lower = [prefix.lower() for prefix in prefixes]

    # Check if the string starts with any of the prefixes
    return any(string_lower.startswith(prefix) for prefix in prefixes_lower)





def table_to_latexxqq(table_obj):
    table       = table_obj.body
    table_index = table_obj.index
    caption     = table_obj.title
    num_columns = table_obj.columns
    num_rows    = table.rows

    # Column alignment - left aligned for text, center for numbers (adjust as needed)
    column_alignment = "|".join(["l" if i == 0 else "c" for i in range(num_columns)])

    # Start the LaTeX table format with center environment
    latex_table = "\\begin{center}\n"
    latex_table += "\\begin{tabular}{" + f"|{column_alignment}|" + "}\n\\hline\n"

    # Iterate through each row of the table
    for row in table.rows:
        row_data = [cell.text.strip().replace("&", "\\&") for cell in row.cells]
        latex_table += " & ".join(row_data) + " \\\\\n\\hline\n"

    # End the LaTeX table format
    latex_table += "\\end{tabular}\n"

    # Add the caption at the bottom, outside of the tabular but still inside the center
    latex_table += f"\\textbf{{Tableau {table_index} – {caption}}}\n"
    latex_table += "\\end{center}\n"

    return latex_table


def clean_caption(caption):
    # Regular expression pattern to match table captions and numbers
    pattern = r'^(tableau|table|tab\.|tabelle|tabla|tabel)[\s]*[\d]*[\s]*[:]*[\s]*'
    # Remove the matched patterns and return the cleaned caption
    return re.sub(pattern, '', caption, flags=re.IGNORECASE).strip()




def table_to_latexx(table_obj):
    table = table_obj.body
    table_index = table_obj.index
    caption = table_obj.title
    num_columns = table_obj.columns
    num_rows = len(table.rows)

    # Page width (can be adjusted as needed)
    page_width_cm = 16

    # Calculate column width based on number of columns
    column_width = f"{page_width_cm / num_columns:.2f}cm"
    column_format = "|".join(["p{" + column_width + "}" for _ in range(num_columns)])

    # Include vertical lines at the beginning and end of the table
    column_format = f"|{column_format}|"

    # Start the LaTeX table environment
    latex_table = "\\begin{table}[H]\n"
    latex_table += "\\centering\n"
    latex_table += "\\begin{tabular}{" + f"{column_format}" + "}\n\\hline\n"

    # Iterate through each row of the table
    for row in table.rows:
        row_data = [cell.text.strip().replace("&", "\\&") for cell in row.cells]
        latex_table += " & ".join(row_data) + " \\\\\n\\hline\n"

    # End the tabular environment and add the caption and label
    latex_table += "\\end{tabular}\n"
    latex_table += f"\\caption{{{caption}}}\n"
    latex_table += f"\\label{{table_{table_index}}}\n"
    latex_table += "\\end{table}\n"

    return latex_table


