from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import lorem

# Create a new Word Document
doc = Document()

# Define the sections and subsections
sections = {
    "Data": ["Subsection 1.1", "Subsection 1.2"],
    "Processing": ["Subsection 2.1", "Subsection 2.2"],
    "Training": ["Subsection 3.1", "Subsection 3.2"]
}

# Loop through the sections and subsections to populate the document
for section, subsections in sections.items():
    # Add section title (Heading 1)
    doc.add_heading(section, level=1)

    for subsection in subsections:
        # Add subsection title (Heading 2)
        doc.add_heading(subsection, level=2)

        # Add sub-subsections (Heading 3) and lorem ipsum text
        for i in range(1, 4):
            sub_subsection_title = f"Sub-Subsection {i} under {subsection}"
            doc.add_heading(sub_subsection_title, level=3)
            paragraph = doc.add_paragraph(lorem.sentence())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Save the document
doc.save('structured_document.docx')
