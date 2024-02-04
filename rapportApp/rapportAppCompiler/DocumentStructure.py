# DocumentStructure.py
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from typing import List, Dict, Optional
from docx import Document
from Scripts.utils import caption_validation, table_to_latexx, clean_caption
from Scripts.constants import TABLE_CAPTION_MISSSING
import uuid
import os


class Image:
    def __init__(self, path, placeholder, title="Please add an image title."):
        self.path = path
        self.placeholder = placeholder
        self.title = title


class Table:
    def __init__(self,index = 0, placeholder = 'Table placeholder', title="Please add a table title.", cols = 1, rows =1, body = ""):
        self.index = index
        self.placeholder = placeholder
        self.title = title
        self.columns = cols
        self.rows = rows
        self.body = body


class SectionFive:
    def __init__(self, title: str, content: str, style: str = "Heading 5"):
        self.title: str = title
        self.content: str = content
        self.style: str = style


class SectionFour:
    def __init__(self, title: str, content: str, style: str = "Heading 4"):
        self.title: str = title
        self.content: str = content
        self.style: str = style
        self.sections: List[SectionFive] = []


class SectionThree:
    def __init__(self, title: str, content: str, style: str = "Heading 3"):
        self.title: str = title
        self.content: str = content
        self.style: str = style
        self.sections: List[SectionFour] = []


class SectionTwo:
    def __init__(self, title: str, content: str, style: str = "Heading 2"):
        self.title: str = title
        self.content: str = content
        self.style: str = style
        self.sections: List[SectionThree] = []


class SectionOne:
    def __init__(self, title: str, content: str, style: str = "Heading 1"):
        self.title: str = title
        self.content: str = content
        self.style: str = style
        self.sections: List[SectionTwo] = []


class DocumentBody:
    def __init__(self):
        self.sections: List[SectionOne] = []


class DocumentWrapper:
    """
    This class will hold all data extracted from a docx file.
    """

    def __init__(
        self,
        logo: Optional[str] = None,
        title: Optional[str] = None,
        subtitle: Optional[str] = None,
        authors: Optional[List[str]] = None,
        course: Optional[str] = None,
        program: Optional[str] = None,
        department: Optional[str] = None,
        supervisor: Optional[str] = None,
        professor: Optional[str] = None,
        date: Optional[str] = None,
        body: Optional[DocumentBody] = None,
        path_to_file: str = None,
    ):
        self.logo = logo
        self.title = title
        self.subtitle = subtitle
        self.authors = authors if authors is not None else []
        self.course = course
        self.program = program
        self.department = department
        self.supervisor = supervisor
        self.date = date
        self.body: DocumentBody = body if body is not None else DocumentBody()
        self.docpath = path_to_file
        self.latex_body = ""
        self.images = []
        self.tables = []
        self.tables_captions = []
        self.erros =[]

        # Call build_structure immediately
        if self.docpath:
            self.store_tables_captions() # Get captions for all tables
            self.add_tables_placeholders() # Store all tables in self.tables and the placeholders in the docuemnt
            if len(self.tables) > 0 :
                self.replace_tables_placeholders() 
                self.clean_tables_stuff()
            self.extract_and_replace_images()
            self.build_structure()
            self.generate_latex_body()
            self.save_latex_body_to_file()

    def __str__(self):
        """
        prints document attributes:
        |-- Title
        |-- Logo
        |-- Subtitle
        |-- Authors
        |-- Course
        |-- Program
        |-- Department
        |-- Supervisor
        |-- Date
        """
        attrs = []
        for key, value in self.__dict__.items():
            # Format the string representation based on whether the attribute has a value
            attr_str = (
                f"|-- {key.capitalize()}: {value}"
                if value is not None
                else f"|-- {key.capitalize()}: ''"
            )
            attrs.append(attr_str)
        return "\n".join(attrs)

    def printDoc(self, sections=None, indent_level=0, onlyStructure=False):
        """
        prints document attributes and sections:
        """
        if sections is None:
            sections = self.body.sections
            print("Document Information:")
            print(f"|-- Title: {self.title}")
            print(f"|-- Logo: {self.logo}")
            print(f"|-- Subtitle: {self.subtitle}")
            print(f"|-- Authors: {', '.join(self.authors) if self.authors else 'None'}")
            print(f"|-- Course: {self.course}")
            print(f"|-- Program: {self.program}")
            print(f"|-- Department: {self.department}")
            print(f"|-- Supervisor: {self.supervisor}")
            print(f"|-- Date: {self.date}")
            print("\nDocument Content:")

        indent = "    " * indent_level  # Increase indentation for each level
        for section in sections:
            print(f"{indent}|-- Title: {section.title}")
            if not onlyStructure:
                print(f"{indent}|-- Style: {section.style}")
                print(f"{indent}|-- Content: {section.content}")

            # Recursively print subsections
            if hasattr(section, "sections") and section.sections:
                self.printDoc(section.sections, indent_level + 1, onlyStructure)

    def build_structure(self):
        doc = Document(self.docpath)
        # Current section references for each level
        current_sections = [None, None, None, None, None]
        latest_active_lvl = None
        levels = [0, 1, 2, 3, 4]

        for paragraph in doc.paragraphs:
            level = None
            content = paragraph.text.strip()

            if paragraph.style.name == "Heading 1":
                level = 0
                if latest_active_lvl is None:
                    latest_active_lvl = 0
                section = SectionOne(content, "")
            elif paragraph.style.name == "Heading 2":
                level = 1
                latest_active_lvl = 1
                section = SectionTwo(content, "")
            elif paragraph.style.name == "Heading 3":
                level = 2
                latest_active_lvl = 2
                section = SectionThree(content, "")
            elif paragraph.style.name == "Heading 4":
                level = 3
                latest_active_lvl = 3
                section = SectionFour(content, "")
            elif paragraph.style.name == "Heading 5":
                level = 4
                latest_active_lvl = 4
                section = SectionFive(content, "")

            if level is not None:
                current_sections[level] = section
                # Check if there is a parent section and append
                if level > 0 and current_sections[level - 1] is not None:
                    current_sections[level - 1].sections.append(section)
                elif level == 0:
                    self.body.sections.append(section)
            else:
                # Check if there is a current section to append the content
                if (
                    latest_active_lvl in levels
                    and current_sections[latest_active_lvl] is not None
                ):
                    current_sections[latest_active_lvl].content += f"{content}"

    def extract_and_replace_images(self, output_dir="LatexDocs/doc_images"):
        # Extract the base file name from the provided path
        base_filename = os.path.basename(self.docpath)

        # Construct the modified file name with the base file name
        modified_filename = "add_imgs_" + base_filename

        # Check if a file with the modified name already exists and delete it
        if os.path.exists(modified_filename):
            os.remove(modified_filename)

        # Load the document
        doc = Document(self.docpath)

        # Create output directory if it doesn't exist
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Initialize a counter for image naming
        image_counter = 1

        # Iterate through paragraphs to find images
        for i, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if "blip" in run._element.xml:
                    # Found an image
                    image_name = f"image{image_counter}"
                    rel_id = [rel for rel in run._element.xpath(".//a:blip/@r:embed")][
                        0
                    ]
                    image_part = doc.part.related_parts[rel_id]

                    # Generate and save the image filename
                    image_filename = os.path.join(
                        output_dir,
                        image_name + os.path.splitext(image_part.partname)[-1],
                    )
                    with open(image_filename, "wb") as img_file:
                        img_file.write(image_part.blob)

                    # Attempt to capture the caption in the next paragraph
                    caption = ""
                    if i + 1 < len(doc.paragraphs):
                        next_paragraph = doc.paragraphs[i + 1]
                        caption = next_paragraph.text.strip()

                    # Caption checking: image or table
                    caption = caption_validation(caption)
                    # Create a new Image object with the image name and caption
                    image_obj = Image(image_filename, image_name, caption)
                    self.images.append(image_obj)

                    # Generate LaTeX-friendly placeholder
                    latex_placeholder = self.generate_latex_figure(image_obj)

                    # Replace the image in the document with the LaTeX placeholder
                    if i < len(doc.paragraphs) - 1:
                        # Clear the paragraph containing the image and add the placeholder
                        doc.paragraphs[i + 1].clear()
                        doc.paragraphs[i + 1].add_run(latex_placeholder)

                    # Increment the image counter
                    image_counter += 1

        # Save the modified document
        if len(self.images) > 0:
            self.docpath = f"docs/{modified_filename}"
            doc.save(self.docpath)


    def generate_latex_figure(self, image_obj):
        # Get the current working directory
        image_path = image_obj.path
        image_title = image_obj.title
        # Check if the current directory is a part of the image pat
        # Modify the image path to be relative from the LatexDocs directory
        image_path = os.path.join("doc_images", os.path.basename(image_path))
        image_path = image_path.replace("\\", "/")

        # Format the LaTeX figure string
        latex_figure = (
            "\\begin{figure}[H]\n"
            "\\centering\n"
            f"\\includegraphics[width=1\\textwidth]{{{image_path}}}\n"
            f"\\caption{{\\label{{{image_title.replace(' ', '_')}}} {image_title}}}"
            "\\end{figure}"
        )

        return latex_figure

    def generate_latex_body(self, sections=None, level=1):
        if sections is None:
            sections = self.body.sections
            # self.latex_body = "\\begin{document}\n"

        for section in sections:
            if level == 1:
                self.latex_body += f"\\section{{{section.title}}}\n"
            elif level == 2:
                self.latex_body += f"\\subsection{{{section.title}}}\n"
            elif level == 3:
                self.latex_body += f"\\subsubsection{{{section.title}}}\n"
            elif level == 4:
                self.latex_body += f"\\paragraph{{{section.title}}}\n"
            elif level == 5:
                self.latex_body += f"\\subparagraph{{{section.title}}}\n"

            if section.content:
                self.latex_body += section.content + "\n"

            if hasattr(section, "sections") and section.sections:
                self.generate_latex_body(section.sections, level + 1)

        if sections is self.body.sections:
            pass
            # self.latex_body += "\\end{document}\n" # reccursion to add end doc to file

    def save_latex_body_to_file(self, savedir="temp", filename="document_body.tex"):
        # Join the directory and filename to create the full path
        filepath = os.path.join(savedir, filename)

        # Check if a file with the same name already exists
        if os.path.exists(filepath):
            # Delete the existing file
            os.remove(filepath)
            print(f"Existing file '{filepath}' was removed.")

        # Write the LaTeX content to the file
        with open(filepath, "w", encoding="utf-8") as file:
            file.write(self.latex_body)

        print(f"Latex body Content is saved to {filepath}")


    def add_tables_placeholders(self):
        # Load the document
        doc = Document(self.docpath)

        # Iterate over the tables
        for table_index, table in enumerate(doc.tables):
            # Convert the table to LaTeX format
            # latex_table = table_to_latex(table)
            table_obj = self.table_to_obj(table, table_index + 1)
            self.tables.append(table_obj)
            # Find the table in the document's element tree and get its parent
            table_element = table._element
            parent_element = table_element.getparent()

            # Add a new paragraph with the table placeholder before the original table
            new_paragraph = Document().add_paragraph(table_obj.placeholder)._element
            parent_element.insert(parent_element.index(table_element), new_paragraph)

            # Remove the original table element
            parent_element.remove(table_element)


        # add captions to tables 
        if len(self.tables) == len(self.tables_captions):
            for index in range(len(self.tables)):
                self.tables[index].title = self.tables_captions[index][1]
        else:
            print(TABLE_CAPTION_MISSSING)
            exit()
        # Save the modified document
        if len(self.tables) > 0:
            modified_docpath = f"modified_{os.path.basename(self.docpath)}"
            modified_docpath = self.docpath.replace(".docx", "modified_.docx")
            self.docpath = modified_docpath
            doc.save(modified_docpath)


    def table_to_obj(self, table, table_index):
        # Placeholder based on the table index
        placeholder = f"xxxtable{table_index}xxx"

        # Calculate the number of columns and rows
        num_cols = len(table.columns)
        num_rows = len(table.rows)

        # Store the table object (you may want to store its content or a representation of it)
        table_body = table  # This is a direct reference to the table object. 
                            # Depending on your requirements, you might want to convert 
                            # this to a more suitable format like a nested list or 
                            # a string representation.

        # Create a Table object
        table_obj = Table(index= table_index, placeholder= placeholder, cols = num_cols,rows= num_rows, body = table_body)

        return table_obj

    def format_tables(self):
        pass


 

    def replace_tables_placeholders(self):
        # Ensure the document path is valid
        if not self.docpath or not os.path.exists(self.docpath):
            print("Document path is invalid or file does not exist.")
            return

        # Load the document using python-docx
        doc = Document(self.docpath)

        # Iterate through the paragraphs to find and replace placeholders
        for table in self.tables:
            for paragraph in doc.paragraphs:
                if table.placeholder in paragraph.text:
                    # Found the placeholder, replace it with the LaTeX string of the table
                    latex_table = table_to_latexx(table)
                    paragraph.text = paragraph.text.replace(table.placeholder, latex_table)

        # Save the modified document
        modified_docpath = self.docpath.replace('.docx', '_tables_replaced.docx')
        self.docpath = modified_docpath
        doc.save(modified_docpath)


    def store_tables_captions(self):
        # Load the document
        doc = Document(self.docpath)

        # Define a broader list of elements that might be used for table captions
        elements_to_search = ['table', 'tableau', 'tab.', 'tabelle', 'tabla', 'tabel']  

        table_index = 1 
        # Loop through each paragraph and check if any element is at the beginning
        for para in doc.paragraphs:
            if para.text:  # Checks if para.text is not None or empty
                # Strip leading/trailing spaces and convert to lower case
                text = para.text.strip().lower()
                if any(text.startswith(element) for element in elements_to_search):
                    caption = clean_caption(para.text)
                    self.tables_captions.append((table_index, caption))
                    table_index += 1


    def clean_tables_stuff(self):
        print(self.tables_captions)
        for tab in self.tables:
            print(tab.index, tab.title , tab.placeholder)

